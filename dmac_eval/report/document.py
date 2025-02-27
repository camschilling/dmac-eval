"""Builds a word document report."""
import datetime

import pandas as pd
from docx import Document
from docx.shared import Pt


def report_preamble(
    report_month_start: datetime.date,
    report_month_end: datetime.date,
) -> Document:
    """
    Begins the word document for the report.

    :param report_month_start: beginning time period the report data includes
    :type report_month_start: datetime.date
    :param report_month_end: end time period the report data includes
    :type report_month_end: datetime.date
    :return: opened word document
    :rtype: Document
    """
    # Create a new Document
    doc = Document()

    # Add a title
    doc.add_heading("DMAC Headcount Report", level=1)

    txt = f"""With reported data from {report_month_start.strftime('%m/%Y')} to
{report_month_end.strftime('%m/%Y')}"""

    doc.add_paragraph(txt.replace("\n", " "))

    # Add a paragraph
    doc.add_paragraph("This report will identify:")

    numbered_items = [
        "Onboarded contractors not billing hours across all projects this month.",
        "Onboarded contractors not billing hours on at least one, but not all, projects this month.",
    ]

    # Add each item as a numbered point
    for item in numbered_items:
        doc.add_paragraph(item, style="ListNumber")

    return doc


def add_df_to_doc(
    doc: Document,
    df: pd.DataFrame,
    title: str,
    font_size: int = 8,
) -> Document:
    """
    Adds a pandas dataframe to a Word document.

    Also includes a header.

    :param doc: opened docx Word document
    :type doc: Document
    :param df: dataframe holding data you'd like to write.
    :type df: pd.DataFrame
    :param title: Title to be included of the table.
    :type title: str
    :param font_size: Font size to have the values in the table be, defaults to 8
    :type font_size: int, optional
    :return: docx Word document, still open
    :rtype: Document
    """
    # Add a title
    doc.add_heading(title, level=4)

    # Add a table with the same number of rows and columns as the DataFrame
    table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])

    # Populate the header row
    for j in range(df.shape[1]):
        cell = table.cell(0, j)
        cell.text = df.columns[j]

    # Populate the rest of the table with data from DataFrame
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            cell = table.cell(
                i + 1,
                j,
            )  # Start filling from row index 1 (after headers)
            cell.text = str(df.iat[i, j])  # Convert to string to avoid type issues  # noqa: PD009

    # Set font size for all cells after filling them with data
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                # Ensure there's at least one run to modify
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                run.font.size = Pt(font_size)  # Set font size to 8 points

    return doc


def save_doc(doc: Document, file_name: str) -> None:
    """
    Saves a Word document to the specified file_name.

    :param doc: a docx Document
    :type doc: docx.Document
    :param file_name: full name for the file (including .docx)
    :type file_name: str
    """
    doc.save(file_name)
